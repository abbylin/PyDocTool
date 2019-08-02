# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'documentTools.py'
#
# Created by: PyQt5 UI code generator 5.9.2
#
# WARNING! All changes made in this file will be lost!

import docx
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from googletrans import Translator

from PyQt5.QtCore import *

class Translater(QThread):
    progressSignal = pyqtSignal(float, str)  # 更新进度条值，提示文字
    finishSignal = pyqtSignal(bool)  #完成处理
    sourceFileName = ''
    targetFileName = ''

    def _init_(self, src='', target=''):
        super(Translater, self).__init__()

    def run(self):
        file = docx.Document(self.sourceFileName)
        totalParagraphs = len(file.paragraphs)

        finalDoc = Document()
        finalDoc.styles['Normal'].font.name = 'Times New Roman'
        finalDoc.styles['Normal'].font.size = Pt(14)
        finalDoc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        if os.path.exists(self.targetFileName):
            os.remove(self.targetFileName)

        translator = Translator()
        for i in range(totalParagraphs):
            # text = "正在处理第"+str(i+1)+"段"+"(共计"+str(totalParagraphs)+"段)"
            # self.progressSignal.emit((i+1)/totalParagraphs, text)
            result = translator.translate(file.paragraphs[i].text, dest="zh-CN")
            result.text.replace("（", "(")
            result.text.replace("）", ")")
            finalDoc.add_paragraph(result.text)
            print("正在处理第" + str(i) + "段")

        finalDoc.save(self.targetFileName)
        self.finishSignal.emit(True)
